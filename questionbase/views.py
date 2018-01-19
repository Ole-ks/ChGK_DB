from django.shortcuts import redirect, render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash
from django.contrib import messages
from django.utils import timezone
from django.http import JsonResponse
from django.http import HttpResponse
from django.db.models import Max
from django.conf import settings
from .models import Question
from .models import Package
from .models import PackageDetail
from .forms import QuestForm
from .forms import PkgForm
from .filters import QFilter
import mimetypes
import os
import shutil
import time
import zipfile

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


@login_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)  # Important!
            messages.success(request, 'Your password was successfully updated!')
            return redirect('search')
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'accounts/change_password.html', {'form': form})


@login_required
def quest_new(request):
    if request.method == 'POST':
        form = QuestForm(request.POST, request.FILES)
        if form.is_valid():
            quest = form.save(commit=False)
            quest.q_has_media = quest.has_media()
            quest.q_has_img = quest.has_img()
            quest.q_has_video = quest.has_video()
            quest.q_has_audio = quest.has_audio()
            quest.created_date = timezone.now()
            quest.save()
            if 'submit_new' in request.POST:
                return redirect('quest_new')
            elif 'submit_edit' in request.POST:
                return redirect('quest_edit', pk=quest.pk)
            else:
                pfrom = request.GET.get('from', None)
                if pfrom:
                    return redirect(pfrom)
                return redirect('search')
    else:
        form = QuestForm(request.POST)
    return render(request, 'questionbase/quest_edit.html', {'form': form})


@login_required
def quest_edit(request, pk):
    quest = get_object_or_404(Question, pk=pk)
    if request.method == "POST":
        form = QuestForm(request.POST, request.FILES, instance=quest)
        if form.is_valid():
            quest = form.save(commit=False)
            quest.q_has_media = quest.has_media()
            quest.q_has_img = quest.has_img()
            quest.q_has_video = quest.has_video()
            quest.q_has_audio = quest.has_audio()
            quest.save()
            if 'submit_new' in request.POST:
                return redirect('quest_new')
            elif 'submit_edit' in request.POST:
                return redirect('quest_edit', pk=quest.pk)
            else:
                pfrom = request.GET.get('from', None)
                if pfrom:
                    return redirect(pfrom + '#q' + pk)
                return redirect('search')
    else:
        form = QuestForm(instance=quest)
    return render(request, 'questionbase/quest_edit.html', {'form': form, 'questid': pk})


@login_required
def search(request):
    quest_filter = QFilter(request.GET, queryset=Question.objects.filter(is_deleted=False).order_by('-created_date'))
    questions = quest_filter.qs

    if 'find-btn' in request.GET:
        request_dict = dict(request.GET)
        del request_dict['find-btn']
        request_dict_url = ''
        for key, value in request_dict.items():
            request_dict_url = request_dict_url + '&' + str(key) + '=' + str(value)[2:-2]
        request.session['search_query'] = request_dict_url
    if 'cancel-btn' in request.GET:
        request.session['search_query'] = ''

    # Paginator parameter
    q_per_page = int(request.GET.get('cou', False))
    if q_per_page in [1, 10, 20, 50]:
        request.session['q_per_page'] = q_per_page
    if 'q_per_page' in request.session:
        paginator = Paginator(questions, request.session['q_per_page'])
    else:
        paginator = Paginator(questions, 10)
    page = request.GET.get('page')
    try:
        questions = paginator.page(page)
    except PageNotAnInteger:
        # If page is not an integer, deliver first page.
        questions = paginator.page(1)
    except EmptyPage:
        # If page is out of range (e.g. 9999), deliver last page of results.
        questions = paginator.page(paginator.num_pages)

    q_in_pkg_list = []
    if 'work_pkg_id' in request.session:
        q_in_pkg_list = PackageDetail.objects.filter(pkg_id=request.session['work_pkg_id']).values_list('quest_id', flat=True)

    return render(request, 'questionbase/quest_search.html', {'filter': quest_filter, 'questions': questions, 'current_path': request.get_full_path(), 'q_in_pkg_list': q_in_pkg_list})


@login_required
def pkg_new(request):
    if request.method == 'POST':
        form = PkgForm(request.POST)
        if form.is_valid():
            pkg = form.save(commit=False)
            pkg.created_date = timezone.now()
            pkg.save()
            request.session['work_pkg_id'] = pkg.id
            request.session['work_pkg_name'] = pkg.name
            return redirect('search')
    else:
        form = PkgForm(request.POST)
    return render(request, 'questionbase/pkg_new.html', {'form': form, 'current_path': request.get_full_path()})


@login_required
def pkg_list(request):
    pkgs = Package.objects.order_by('-created_date')
    return render(request, 'questionbase/pkg_list.html', {'pkgs': pkgs, 'current_path': request.get_full_path()})


@login_required
def pkg_detail(request, pk):
    pkg = get_object_or_404(Package, pk=pk)
    questions = Question.objects.filter(question_in_pkg__pkg_id=pk).order_by('question_in_pkg__num_in_pkg')
    #questions = PackageDetail.objects.filter(pkg_id=pk).select_related('quest_id')
    request.session['work_pkg_id'] = pkg.id
    request.session['work_pkg_name'] = pkg.name
    return render(request, 'questionbase/pkg_detail.html', {'questions': questions, 'current_path': request.get_full_path(), 'is_done': pkg.is_done})


@login_required
def search_w_a(request):
    request.session['with_answers'] = 'True'
    pfrom = request.GET.get('from', None)
    if pfrom:
        return redirect(pfrom)
    return redirect('search')


@login_required
def search_n_a(request):
    request.session['with_answers'] = 'False'
    pfrom = request.GET.get('from', None)
    if pfrom:
        return redirect(pfrom)
    return redirect('search')


@login_required
def quest_remove(request, pk):
    quest = get_object_or_404(Question, pk=pk)
    quest.is_deleted = True
    quest.save()
    return redirect('search')


@login_required
def quest_restore(request, pk):
    quest = get_object_or_404(Question, pk=pk)
    quest.is_deleted = False
    quest.save()
    return redirect('search')


@login_required
def delete_forever(request, pk):
    quest = get_object_or_404(Question, pk=pk)
    for file in [quest.qlink1, quest.qlink2, quest.qlink3, quest.alink1, quest.alink2, quest.alink3]:
        try:
            file.delete()
            quest.save()
        except:
            print('-------something went wrong++++++++++++++++++++++++++++++++++++++++++++')
    quest.delete()
    return redirect('deleted')


@login_required
def quest_remove_from_pkg(request, pk):
    quest = get_object_or_404(Question, pk=pk)
    pkg = get_object_or_404(Package, pk=request.session['work_pkg_id'])
    pkg_det = get_object_or_404(PackageDetail, pkg_id=pkg, quest_id=quest)
    pkg_det.delete()
    pk = request.session['work_pkg_id']
    return redirect('pkg_detail', pk)


@login_required
def deleted(request):
    questions = Question.objects.filter(is_deleted=True).order_by('-created_date')

    # Paginator parameter
    q_per_page = int(request.GET.get('cou', False))
    if q_per_page in [1, 10, 20, 50]:
        request.session['q_per_page'] = q_per_page

    if 'q_per_page' in request.session:
        paginator = Paginator(questions, request.session['q_per_page'])
    else:
        paginator = Paginator(questions, 10)
    page = request.GET.get('page')
    try:
        questions = paginator.page(page)
    except PageNotAnInteger:
        # If page is not an integer, deliver first page.
        questions = paginator.page(1)
    except EmptyPage:
        # If page is out of range (e.g. 9999), deliver last page of results.
        questions = paginator.page(paginator.num_pages)

    return render(request, 'questionbase/deleted.html', {'questions': questions, 'current_path': request.get_full_path()})


@login_required
def add_quest_to_pkg(request):
    q_id = request.GET.get('q_id', None)
    if q_id:
        quest = get_object_or_404(Question, pk=q_id)
        pkg = get_object_or_404(Package, pk=request.session['work_pkg_id'])
        last_num = PackageDetail.objects.filter(pkg_id=request.session['work_pkg_id']).aggregate(Max('num_in_pkg'))
        if not last_num['num_in_pkg__max']:
            new_num = 1
        else:
            new_num = int(last_num['num_in_pkg__max']) + 1
        p = PackageDetail(quest_id=quest, pkg_id=pkg, num_in_pkg=new_num)
        p.save()
        data = {
            'cou': str(new_num)
        }
    return JsonResponse(data)


@login_required
def show_pkg_for_quest(request):
    b_id = request.GET.get('b_id', None)
    q_id = b_id[6:]

    pkg_list = Package.objects.filter(pkg_detail__quest_id=q_id).values_list('name', flat=True)

    data = {
        'q_id': q_id,
        'pkg_list': list(pkg_list)
    }
    return JsonResponse(data)


@login_required
def pkg_ready(request):
    pkg = get_object_or_404(Package, pk=request.session['work_pkg_id'])
    pkg.is_done = True
    pkg.save()
    pk = request.session['work_pkg_id']
    return redirect('pkg_detail', pk)


@login_required
def pkg_not_ready(request):
    pkg = get_object_or_404(Package, pk=request.session['work_pkg_id'])
    pkg.is_done = False
    pkg.save()
    pk = request.session['work_pkg_id']
    return redirect('pkg_detail', pk)


@login_required
def write_pdf_view(request):
    pk = request.session['work_pkg_id']
    pkg = Package.objects.get(pk=pk)
    questions = Question.objects.filter(question_in_pkg__pkg_id=pk).order_by('question_in_pkg__num_in_pkg')

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    document.add_paragraph(pkg.name)
    for quest in questions:
        table = document.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False

        for link, type, i in quest.ext_q_list():
            if type == 'img':
                table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run('/картинка/ ')
            if type == 'video':
                table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run('/видео/ ')
            if type == 'audio':
                table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run('/аудио/ ')

        table.cell(0,0).merge(table.cell(0,1)).paragraphs[0].add_run('Вопрос. ').bold = True
        table.cell(0,0).merge(table.cell(0,1)).paragraphs[0].add_run(quest.qtext).bold = False

        for link, type, i in quest.ext_a_list():
            if type == 'img':
                table.cell(1, 0).paragraphs[0].add_run('/картинка/ ')
            if type == 'video':
                table.cell(1, 0).paragraphs[0].add_run('/видео/ ')
            if type == 'audio':
                table.cell(1, 0).paragraphs[0].add_run('/аудио/ ')

        table.cell(1,0).paragraphs[0].add_run('Ответ: ').bold = True
        table.cell(1,0).paragraphs[0].add_run(quest.answer).bold = False
        table.cell(1,1).paragraphs[0].add_run('Автор: ').bold = True
        table.cell(1,1).paragraphs[0].add_run(quest.author).bold = False
        table.cell(2,0).paragraphs[0].add_run('Комментарий: ').bold = True
        table.cell(2, 0).paragraphs[0].add_run(quest.comment).bold = False
        table.cell(2,1).paragraphs[0].add_run('Источник: ').bold = True
        table.cell(2, 1).paragraphs[0].add_run(quest.source).bold = False

        document.add_paragraph('')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response


@login_required
def download_files(request):
    pk = request.session['work_pkg_id']
    questions = Question.objects.filter(question_in_pkg__pkg_id=pk).order_by('question_in_pkg__num_in_pkg')

    #collect [filename, folder, filename_new] to copy
    list = []
    i = 0
    for quest in questions:
        i += 1
        if quest.qlink1:
            list.append([str(quest.qlink1.name).split('/')[-1], quest.qlink1.url, str(i) + '_' + str(quest.qlink1.name).replace('/', '_')])
        if quest.qlink2:
            list.append([str(quest.qlink2.name).split('/')[-1], quest.qlink2.url, str(i) + '_' + str(quest.qlink2.name).replace('/', '_')])
        if quest.qlink3:
            list.append([str(quest.qlink3.name).split('/')[-1], quest.qlink3.url, str(i) + '_' + str(quest.qlink3.name).replace('/', '_')])
        if quest.alink1:
            list.append([str(quest.alink1.name).split('/')[-1], quest.alink1.url, str(i) + '_' + str(quest.alink1.name).replace('/', '_')])
        if quest.alink2:
            list.append([str(quest.alink2.name).split('/')[-1], quest.alink2.url, str(i) + '_' + str(quest.alink2.name).replace('/', '_')])
        if quest.alink3:
            list.append([str(quest.alink3.name).split('/')[-1], quest.alink3.url, str(i) + '_' + str(quest.alink3.name).replace('/', '_')])

    #create new temp dir
    folder_new = os.path.join(settings.BASE_DIR, 'questionbase', 'media', 'temp_' + str(pk) + '_' + str(int(time.time())), 'img')
    os.makedirs(folder_new)

    #copy files to temp dir
    for item in list:
        filepath = os.path.join(settings.BASE_DIR, 'questionbase', str(item[1])[1:])
        filepath_new = os.path.join(folder_new, item[2])
        shutil.copy2(filepath, filepath_new)

    # make archive
    archive_path = os.path.join(settings.BASE_DIR, 'questionbase', 'media', 'temp_' + str(pk) + '_' + str(int(time.time())), 'img')
    shutil.make_archive(archive_path, 'zip', folder_new)

    fp = open(archive_path + '.zip', 'rb')
    response = HttpResponse(fp.read())
    fp.close()
    response['Content-Disposition'] = 'attachment; filename="img.zip"'

    return response